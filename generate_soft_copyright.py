# -*- coding: utf-8 -*-
"""
软著代码文档生成器 - 一鉴到底 (YiJianDaoDi)
生成前30页 + 后60页 源代码 Word 文档
逻辑连续，不分页断层
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime

# ── 项目根目录 ──
PROJECT_ROOT = Path(r"C:\MsSafeData\Desktop\yijiandaodi")

# ── 输出路径 ──
OUTPUT_PATH = PROJECT_ROOT / "一鉴到底_源代码_软著.docx"

# ── 分页参数 ──
LINES_PER_PAGE = 50       # 每页行数（软著标准）
FIRST_PAGES = 30          # 前30页
LAST_PAGES = 60           # 后60页
FIRST_LINES = FIRST_PAGES * LINES_PER_PAGE   # 1500 行
LAST_LINES = LAST_PAGES * LINES_PER_PAGE      # 3000 行

# ── 文件读取顺序（逻辑连续）──
# 排除：node_modules, __pycache__, .git, dist, build, venv, *.pyc, migrations

FILE_ORDER = [
    # ══════════════════════════════
    # 第一部分：前端 - 入口与配置
    # ══════════════════════════════
    "frontend/package.json",
    "frontend/tsconfig.json",
    "frontend/vite.config.ts",
    "frontend/index.html",
    "frontend/src/main.tsx",
    "frontend/src/App.tsx",
    "frontend/src/router/index.tsx",

    # ══════════════════════════════
    # 第二部分：前端 - API 层
    # ══════════════════════════════
    "frontend/src/utils/request.ts",
    "frontend/src/api/agentApi.ts",
    "frontend/src/api/skillConfigApi.ts",
    "frontend/src/api/userApi.ts",
    "frontend/src/api/authApi.ts",
    "frontend/src/api/paymentApi.ts",
    "frontend/src/api/orderApi.ts",

    # ══════════════════════════════
    # 第三部分：前端 - 工具层
    # ══════════════════════════════
    "frontend/src/utils/contextManager.ts",
    "frontend/src/utils/token.ts",
    "frontend/src/utils/storage.ts",
    "frontend/src/utils/format.ts",

    # ══════════════════════════════
    # 第四部分：前端 - 数据层
    # ══════════════════════════════
    "frontend/src/data/skillMatrix.ts",
    "frontend/src/data/agentRoles.ts",
    "frontend/src/store/index.ts",
    "frontend/src/store/useAuthStore.ts",

    # ══════════════════════════════
    # 第五部分：前端 - 核心页面（Agent对话）
    # ══════════════════════════════
    "frontend/src/pages/Home/components/AIChatCenter.tsx",
    "frontend/src/pages/YijiandaodiSkill.tsx",
    "frontend/src/pages/executioncenter/index.tsx",
    "frontend/src/pages/AgentSkillDetail.tsx",

    # ══════════════════════════════
    # 第六部分：前端 - 布局组件
    # ══════════════════════════════
    "frontend/src/layouts/MainLayout.tsx",
    "frontend/src/layouts/BrandLayout.tsx",
    "frontend/src/components/Header/NavBar.tsx",
    "frontend/src/components/Footer/Footer.tsx",

    # ══════════════════════════════
    # 第七部分：后端 - 配置与入口
    # ══════════════════════════════
    "backend/manage.py",
    "backend/fangdudu_backend/__init__.py",
    "backend/fangdudu_backend/settings.py",
    "backend/fangdudu_backend/urls.py",
    "backend/fangdudu_backend/wsgi.py",
    "backend/fangdudu_backend/asgi.py",

    # ══════════════════════════════
    # 第八部分：后端 - URL 路由
    # ══════════════════════════════
    "backend/auth_app/agent_urls.py",
    "backend/auth_app/urls.py",
    "backend/content_app/urls.py",
    "backend/p2p_app/urls.py",

    # ══════════════════════════════
    # 第九部分：后端 - 数据模型
    # ══════════════════════════════
    "backend/auth_app/models.py",
    "backend/auth_app/agent_models.py",
    "backend/content_app/models.py",
    "backend/p2p_app/models.py",

    # ══════════════════════════════
    # 第十部分：后端 - 视图（核心业务）
    # ══════════════════════════════
    "backend/auth_app/agent_views.py",
    "backend/auth_app/views.py",
    "backend/content_app/views.py",
    "backend/p2p_app/views.py",

    # ══════════════════════════════
    # 第十一部分：后端 - AI 服务
    # ══════════════════════════════
    "backend/content_app/deepseek_service.py",
    "backend/content_app/pyodide_service.py",

    # ══════════════════════════════
    # 第十二部分：后端 - 中间件 & 序列化器
    # ══════════════════════════════
    "backend/fangdudu_backend/security_middleware.py",
    "backend/auth_app/serializers.py",
    "backend/content_app/serializers.py",
    "backend/p2p_app/serializers.py",

    # ══════════════════════════════
    # 第十三部分：后端 - P2P 分布式核心
    # ══════════════════════════════
    "backend/p2p_app/node_discovery.py",
    "backend/p2p_app/task_scheduler.py",
    "backend/p2p_app/result_aggregator.py",
    "backend/p2p_app/eihm_router.py",
    "backend/p2p_app/sandbox_executor.py",
]


def read_file_lines(filepath: Path) -> list[str]:
    """读取文件所有行，跳过空文件和二进制文件"""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.readlines()
    except Exception:
        return []


def collect_all_lines() -> list[tuple[str, int, str]]:
    """
    收集所有文件的代码行
    返回: [(file_path, line_number, content), ...]
    """
    all_lines = []
    for rel_path in FILE_ORDER:
        full_path = PROJECT_ROOT / rel_path
        if not full_path.exists():
            print(f"  [跳过] 不存在: {rel_path}")
            continue

        lines = read_file_lines(full_path)
        if not lines:
            continue

        for i, line in enumerate(lines, 1):
            all_lines.append((rel_path, i, line.rstrip('\n\r')))

        print(f"  [OK] {rel_path} ({len(lines)} 行)")

    return all_lines


def generate_docx(all_lines: list[tuple[str, int, str]], output_path: Path):
    """使用 python-docx 生成 Word 文档"""

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("正在安装 python-docx...")
        os.system("pip install python-docx -q")
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)     # A4 宽
    section.page_height = Cm(29.7)  # A4 高
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 样式设置 ──
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    total_lines = len(all_lines)
    print(f"\n总代码行数: {total_lines}")

    # ── 封面信息 ──
    title = doc.add_heading('', level=0)
    run = title.add_run('一鉴底底 (YiJianDaoDi)')
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(22, 93, 255)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run('计算机软件著作权登记申请 · 源代码文档')
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = (
        f"前 {FIRST_PAGES} 页 + 后 {LAST_PAGES} 页  |  "
        f"共 {FIRST_PAGES + LAST_PAGES} 页  |  "
        f"每页 {LINES_PER_PAGE} 行  |  "
        f"总代码量 {total_lines:,} 行\n"
        f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    info_para.add_run(info_text).font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # ── 目录索引 ──
    toc_title = doc.add_heading('源代码文件目录', level=1)
    toc_title.runs[0].font.size = Pt(16)

    seen_files = set()
    for fp, _, _ in all_lines:
        if fp not in seen_files:
            seen_files.add(fp)
            p = doc.add_paragraph()
            p.add_run(f"  • {fp}").font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 第一部分：前 30 页（从头开始）
    # ══════════════════════════════════════
    part1_title = doc.add_heading(
        f'第一部分：源代码前 {FIRST_PAGES} 页（第 1 ~ {FIRST_LINES} 行）',
        level=1
    )
    part1_title.runs[0].font.size = Pt(14)
    part1_title.runs[0].font.color.rgb = RGBColor(22, 93, 255)

    current_file = None
    part1_end = min(FIRST_LINES, total_lines)

    for idx in range(part1_end):
        file_path, line_no, content = all_lines[idx]

        # 文件切换时加标题
        if file_path != current_file:
            current_file = file_path
            fh = doc.add_heading(file_path, level=2)
            fh.runs[0].font.size = Pt(11)
            fh.runs[0].font.color.rgb = RGBColor(80, 80, 80)

        # 代码行
        p = doc.add_paragraph()
        code_text = f"{line_no:>5}  │{content}"
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

        # 每 LINES_PER_PAGE 行不强制分页，保持逻辑连续
        # （Word 会自然分页）

    doc.add_page_break()

    # ══════════════════════════════════════
    # 第二部分：后 60 页（从尾部回溯）
    # ══════════════════════════════════════
    part2_start = max(0, total_lines - LAST_LINES)

    part2_title = doc.add_heading(
        f'第二部分：源代码后 {LAST_PAGES} 页'
        f'（第 {part2_start + 1} ~ {total_lines} 行）',
        level=1
    )
    part2_title.runs[0].font.size = Pt(14)
    part2_title.runs[0].font.color.rgb = RGBColor(14, 165, 233)

    current_file = None

    for idx in range(part2_start, total_lines):
        file_path, line_no, content = all_lines[idx]

        # 文件切换时加标题
        if file_path != current_file:
            current_file = file_path
            fh = doc.add_heading(file_path, level=2)
            fh.runs[0].font.size = Pt(11)
            fh.runs[0].font.color.rgb = RGBColor(80, 80, 80)

        # 代码行
        p = doc.add_paragraph()
        code_text = f"{line_no:>5}  │{content}"
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

    # ── 结尾声明 ──
    doc.add_paragraph()
    doc.add_paragraph()
    decl = doc.add_paragraph()
    decl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decl_run = decl.add_run(
        '— 以上为「一鉴底底」系统全部源代码的前30页及后60页，'
        '代码逻辑连续完整，无删减或拼接 —'
    )
    decl_run.font.size = Pt(10)
    decl_run.font.color.rgb = RGBColor(150, 150, 150)
    decl_run.font.italic = True

    # ── 保存 ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    file_size_mb = output_path.stat().st_size / (1024 * 1024)

    print(f"\n{'='*60}")
    print(f"✅ 软著代码文档已生成!")
    print(f"   📄 输出: {output_path}")
    print(f"   📊 大小: {file_size_mb:.2f} MB")
    print(f"   📝 总行数: {total_lines:,}")
    print(f"   📑 前{FIRST_PAGES}页: 行 1 ~ {part1_end}")
    print(f"   📑 后{LAST_PAGES}页: 行 {part2_start+1} ~ {total_lines}")
    print(f"   📁 文件数: {len(seen_files)}")
    print(f"{'='*60}")


def main():
    print("=" * 60)
    print("  一鉴到底 (YiJianDaoDi) 软著代码文档生成器")
    print(f"  前{FIRST_PAGES}页 + 后{LAST_PAGES}页 | 每页{LINES_PER_PAGE}行")
    print("=" * 60)
    print("\n📂 正在扫描源码文件...\n")

    all_lines = collect_all_lines()

    if len(all_lines) < FIRST_LINES + LAST_LINES:
        print(f"\n⚠️  警告: 总行数({len(all_lines)}) 少于要求行数"
              f"({FIRST_LINES + LAST_LINES})")
        print("   将输出所有可用代码。")

    generate_docx(all_lines, OUTPUT_PATH)


if __name__ == '__main__':
    main()
